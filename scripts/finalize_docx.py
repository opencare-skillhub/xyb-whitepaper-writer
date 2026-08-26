#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
finalize_docx.py — 定稿：替换图片 marker + 每章分页 + 校验

读取 finalize.json：
{
  "docx": "stage3/target.docx",
  "images_dir": "stage2",
  "markers": {
     "INSERT_IMAGE_COVER_LOGO": {"file": "logo_mascot.png", "width_cm": 3.2},
     "INSERT_IMAGE_COMPARE": {"file": "compare.png", "width_cm": 15.5},
     "INSERT_IMAGE_COMMUNITY_LOGO": {"file": "logo_mascot.png", "width_cm": 6.0}
  },
  "skip_first_h1": true
}

动作：
  1) 把 docx 中的 {{INSERT_IMAGE_x}} 文本（段落或表格单元格）替换为真实图片（指定 Cm 宽、居中）
  2) 每个 H1 设 page_break_before（skip_first_h1=true 时跳过封面标题，避免空白页）
  3) 校验并报告：图片数、marker 残留、H1 数、分页数

依赖：python-docx；图片相对 images_dir 解析。
"""
import os, sys, json, argparse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.oxml.ns import qn


def replace_markers_in_paragraph(par, markers, images_dir):
    replaced = 0
    for m, spec in markers.items():
        # 转换器可能把 {{X}} 规范成 {X} 甚至保留 {{X}}，逐一尝试
        candidates = ["{{" + m + "}}", "{" + m + "}", m]
        full = "".join(r.text or "" for r in par.runs)
        token = next((c for c in candidates if c in full), None)
        if token is None:
            continue
        img = os.path.join(images_dir, spec["file"])
        w = spec.get("width_cm", 12.0)
        if not os.path.exists(img):
            print(f"  [WARN] marker {m} 图片缺失: {img}")
            continue
        # 定位含 token 的 run
        target = None
        for r in par.runs:
            if token in (r.text or ""):
                target = r
                break
        if target is None:
            # token 跨 run 拆分：清空全部 run，直接加图
            for r in list(par.runs):
                r._r.getparent().remove(r._r)
            pic_run = par.add_run()
            pic_run.add_picture(img, width=Cm(w))
            pic_run.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            target.text = (target.text or "").replace(token, "")
            pic_run = par.add_run()
            pic_run.add_picture(img, width=Cm(w))
            pic_run.alignment = WD_ALIGN_PARAGRAPH.CENTER
            target._r.addprevious(pic_run._r)
        replaced += 1
        print(f"[img] replaced {m} -> {spec['file']} @ {w}cm")
        break  # 一个段落最多一个 marker
    return replaced


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--json", required=True, help="finalize.json 路径")
    args = ap.parse_args()
    with open(args.json, encoding="utf-8") as f:
        cfg = json.load(f)

    docx_path = cfg["docx"]
    images_dir = cfg.get("images_dir", os.path.dirname(docx_path))
    markers = cfg.get("markers", {})
    skip_first_h1 = bool(cfg.get("skip_first_h1", True))

    if not os.path.exists(docx_path):
        sys.exit(f"ERROR: docx 不存在: {docx_path}（请先运行 html_to_docx convert）")

    doc = Document(docx_path)

    # 1) 替换 marker（段落 + 表格单元格）
    total = 0
    for p in doc.paragraphs:
        total += replace_markers_in_paragraph(p, markers, images_dir)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    total += replace_markers_in_paragraph(p, markers, images_dir)
    print(f"[img] total replaced = {total}")

    # 2) 分页
    breaks = 0
    first_seen = False
    for p in doc.paragraphs:
        if p.style and p.style.name == "Heading 1":
            if skip_first_h1 and not first_seen:
                first_seen = True
                continue
            p.paragraph_format.page_break_before = True
            breaks += 1
    print(f"[page] H1 page_break_before set = {breaks}")

    doc.save(docx_path)

    # 3) 校验
    d2 = Document(docx_path)
    txt = "\n".join(par.text for par in d2.paragraphs)
    for t in d2.tables:
        for row in t.rows:
            for c in row.cells:
                txt += "\n" + c.text
    blips = [el for el in d2.element.body.iter() if el.tag == qn('a:blip')]
    h1 = sum(1 for p in d2.paragraphs if p.style and p.style.name == "Heading 1")
    leftover_marker = sum(1 for m in markers if m in txt)
    print("\n=== VERIFY ===")
    print(f"  images(blips) : {len(blips)}")
    print(f"  marker leftover: {leftover_marker}")
    print(f"  H1 count      : {h1}")
    print(f"  paragraphs    : {len(d2.paragraphs)}  tables: {len(d2.tables)}")
    print(f"  saved         : {docx_path}")
    if leftover_marker > 0:
        print("  [FAIL] 存在未替换的 marker！")
        sys.exit(1)
    print("  [OK] 校验通过")


if __name__ == "__main__":
    main()
